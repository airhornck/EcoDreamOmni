"""
BrandKnowledge — API contract tests.
Tests for existing endpoints only; skip tests for not-yet-implemented endpoints.
"""

import pytest
import uuid
from src.models.user import clear_users


def get_auth_token(client, role: str = "operator"):
    clear_users()
    email = f"bk_{uuid.uuid4().hex[:8]}@ecodream.com"
    response = client.post("/auth/register", json={
        "email": email,
        "password": "SecurePass123!",
        "username": f"bkuser_{uuid.uuid4().hex[:8]}",
        "role": "operator",
    })
    assert response.status_code == 201, f"Register failed: {response.text}"
    return response.json()["access_token"]


# === Existing endpoint tests ===

def test_create_brand_knowledge_entry(client):
    token = get_auth_token(client)
    response = client.post(
        "/brand-knowledge/entries",
        json={
            "entry_type": "PRODUCT_INFO",
            "name": "Test Product",
            "content": "Product description",
            "category": "驱虫",
            "tags": ["test", "product"],
        },
        headers={"Authorization": f"Bearer {token}"},
    )
    assert response.status_code == 201
    data = response.json()
    assert "id" in data
    assert data["name"] == "Test Product"


def test_get_brand_knowledge_entry(client):
    token = get_auth_token(client)
    create_resp = client.post(
        "/brand-knowledge/entries",
        json={
            "entry_type": "CATEGORY_KNOWLEDGE",
            "name": "Category Knowledge",
            "content": "Knowledge content",
            "category": "驱虫",
            "tags": ["knowledge"],
        },
        headers={"Authorization": f"Bearer {token}"},
    )
    entry_id = create_resp.json()["id"]

    response = client.get(f"/brand-knowledge/entries/{entry_id}", headers={"Authorization": f"Bearer {token}"})
    assert response.status_code in (200, 201)
    data = response.json()
    assert data["id"] == entry_id


def test_list_brand_knowledge_entries(client):
    token = get_auth_token(client)
    response = client.get("/brand-knowledge/entries?category=驱虫", headers={"Authorization": f"Bearer {token}"})
    assert response.status_code in (200, 201)
    data = response.json()
    assert "items" in data


def test_update_brand_knowledge_entry(client):
    token = get_auth_token(client)
    create_resp = client.post(
        "/brand-knowledge/entries",
        json={
            "entry_type": "FAQ",
            "name": "FAQ Entry",
            "content": "Original content",
            "category": "营养",
        },
        headers={"Authorization": f"Bearer {token}"},
    )
    entry_id = create_resp.json()["id"]
    original_version = create_resp.json()["version"]

    response = client.put(
        f"/brand-knowledge/entries/{entry_id}",
        json={"content": "Updated content"},
        headers={"Authorization": f"Bearer {token}"},
    )
    assert response.status_code in (200, 201)
    data = response.json()
    assert data["version"] == original_version + 1


# === Not-yet-implemented endpoint stubs (skipped) ===
# The following endpoints are defined in PRD but not yet implemented:
#   POST /brand-knowledge/products
#   POST /brand-knowledge/validate-approval
#   POST /brand-knowledge/check-content
#   POST /brand-knowledge/bulk-import
#   GET  /brand-knowledge/retrieve (RAG semantic search)
#
# When implementing, add tests here and remove the skip markers.


@pytest.mark.skip(reason="Endpoint not yet implemented")
def test_create_product_with_approval_number(client):
    pass


@pytest.mark.skip(reason="Endpoint not yet implemented")
def test_validate_veterinary_approval_number(client):
    pass


@pytest.mark.skip(reason="Endpoint not yet implemented")
def test_detect_prohibited_claims(client):
    pass


@pytest.mark.skip(reason="Endpoint not yet implemented")
def test_check_content_against_knowledge_base(client):
    pass


@pytest.mark.skip(reason="Endpoint not yet implemented")
def test_retrieve_knowledge_for_rag(client):
    pass


def test_bulk_import_csv(client):
    """🔴 CSV bulk import brand knowledge entries."""
    token = get_auth_token(client)
    csv_content = (
        "entry_type,name,content,brand_name,prohibited_claims\n"
        "faq,FAQ1,Answer1,BrandA,claim1;claim2\n"
        "faq,FAQ2,Answer2,BrandB,\n"
    )
    from io import BytesIO

    response = client.post(
        "/brand-knowledge/bulk-import",
        files={"file": ("test.csv", BytesIO(csv_content.encode("utf-8-sig")), "text/csv")},
        headers={"Authorization": f"Bearer {token}"},
    )
    assert response.status_code == 200
    data = response.json()
    assert data["imported_count"] == 2
    assert len(data["errors"]) == 0


def test_bulk_import_pdf(client):
    """🔴 PDF bulk import — one entry per page."""
    token = get_auth_token(client)
    from io import BytesIO
    from PyPDF2 import PdfWriter

    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)
    writer.add_blank_page(width=612, height=792)
    pdf_bytes = BytesIO()
    writer.write(pdf_bytes)
    pdf_bytes.seek(0)

    response = client.post(
        "/brand-knowledge/bulk-import",
        files={"file": ("test.pdf", pdf_bytes, "application/pdf")},
        headers={"Authorization": f"Bearer {token}"},
    )
    assert response.status_code == 200
    data = response.json()
    # Blank pages yield no text, so imported count is 0
    assert data["imported_count"] == 0


def test_bulk_import_docx(client):
    """🔴 Word bulk import — whole document as single entry."""
    token = get_auth_token(client)
    from io import BytesIO
    from docx import Document

    doc = Document()
    doc.add_heading("Product Manual", level=1)
    doc.add_paragraph("This is the product description for BrandX.")
    docx_bytes = BytesIO()
    doc.save(docx_bytes)
    docx_bytes.seek(0)

    response = client.post(
        "/brand-knowledge/bulk-import",
        files={"file": ("test.docx", docx_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        headers={"Authorization": f"Bearer {token}"},
    )
    assert response.status_code == 200
    data = response.json()
    assert data["imported_count"] == 1
    assert data["errors"] == []


def test_bulk_import_xlsx(client):
    """🔴 Excel bulk import — one entry per row."""
    token = get_auth_token(client)
    from io import BytesIO
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.append(["Name", "Description", "Brand"])
    ws.append(["Entry A", "Content A", "BrandA"])
    ws.append(["Entry B", "Content B", "BrandB"])
    xlsx_bytes = BytesIO()
    wb.save(xlsx_bytes)
    xlsx_bytes.seek(0)

    response = client.post(
        "/brand-knowledge/bulk-import",
        files={"file": ("test.xlsx", xlsx_bytes, "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")},
        headers={"Authorization": f"Bearer {token}"},
    )
    assert response.status_code == 200
    data = response.json()
    assert data["imported_count"] == 2
    assert data["errors"] == []


def test_bulk_import_unsupported_format(client):
    """🔴 Unsupported file format rejected."""
    token = get_auth_token(client)
    from io import BytesIO

    response = client.post(
        "/brand-knowledge/bulk-import",
        files={"file": ("test.txt", BytesIO(b"hello"), "text/plain")},
        headers={"Authorization": f"Bearer {token}"},
    )
    assert response.status_code == 400
    assert "Supported formats" in response.json()["detail"]
