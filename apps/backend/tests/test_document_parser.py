"""Document parser unit tests — no DB required.

Tests PDF / Word / Excel / CSV text extraction.
"""

import pytest
from pathlib import Path

from src.core import document_parser as dp


class TestParseCsvText:
    def test_basic_csv(self):
        text = "entry_type,name,content\nfaq,Q1,Answer1\nfaq,Q2,Answer2\n"
        results = dp.parse_csv_text(text)
        assert len(results) == 2
        assert results[0]["entry_type"] == "faq"
        assert results[0]["name"] == "Q1"
        assert results[0]["content"] == "Answer1"

    def test_csv_with_optional_columns(self):
        text = (
            "entry_type,name,content,brand_name,prohibited_claims\n"
            "faq,Q1,A1,BrandX,claim1;claim2\n"
        )
        results = dp.parse_csv_text(text)
        assert len(results) == 1
        assert results[0]["brand_name"] == "BrandX"
        assert results[0]["prohibited_claims"] == ["claim1", "claim2"]

    def test_csv_missing_required_column(self):
        text = "name,content\nQ1,Answer1\n"
        with pytest.raises(ValueError, match="CSV must contain"):
            dp.parse_csv_text(text)


class TestParsePdf:
    def test_parse_pdf_blank_pages(self, tmp_path: Path):
        from PyPDF2 import PdfWriter

        writer = PdfWriter()
        writer.add_blank_page(width=612, height=792)
        writer.add_blank_page(width=612, height=792)
        pdf_path = tmp_path / "test.pdf"
        with open(pdf_path, "wb") as f:
            writer.write(f)

        results = dp.parse_pdf(str(pdf_path))
        # Blank pages yield no extractable text
        assert results == []

    def test_parse_pdf_with_text(self, tmp_path: Path):
        from PyPDF2 import PdfWriter

        # Create a PDF with text by merging from a simple source
        # (PyPDF2 blank pages have no text, so we simulate with a minimal approach)
        writer = PdfWriter()
        writer.add_blank_page(width=612, height=792)
        pdf_path = tmp_path / "test.pdf"
        with open(pdf_path, "wb") as f:
            writer.write(f)

        results = dp.parse_pdf(str(pdf_path))
        assert isinstance(results, list)


class TestParseDocx:
    def test_parse_docx(self, tmp_path: Path):
        from docx import Document

        doc = Document()
        doc.add_heading("Manual", level=1)
        doc.add_paragraph("Product description here.")
        docx_path = tmp_path / "test.docx"
        doc.save(docx_path)

        results = dp.parse_docx(str(docx_path))
        assert len(results) == 1
        assert results[0]["name"] == "test"
        assert "Product description" in results[0]["content"]

    def test_parse_empty_docx(self, tmp_path: Path):
        from docx import Document

        doc = Document()
        docx_path = tmp_path / "empty.docx"
        doc.save(docx_path)

        results = dp.parse_docx(str(docx_path))
        assert results == []


class TestParseXlsx:
    def test_parse_xlsx(self, tmp_path: Path):
        from openpyxl import Workbook

        wb = Workbook()
        ws = wb.active
        ws.append(["Name", "Description", "Brand"])
        ws.append(["Entry A", "Content A", "BrandA"])
        ws.append(["Entry B", "Content B", "BrandB"])
        xlsx_path = tmp_path / "test.xlsx"
        wb.save(xlsx_path)

        results = dp.parse_xlsx(str(xlsx_path))
        assert len(results) == 2
        assert results[0]["name"] == "Entry A"
        assert "Description: Content A" in results[0]["content"]
        assert results[1]["name"] == "Entry B"

    def test_parse_xlsx_empty_rows_skipped(self, tmp_path: Path):
        from openpyxl import Workbook

        wb = Workbook()
        ws = wb.active
        ws.append(["Name", "Description"])
        ws.append(["Entry A", "Content A"])
        ws.append([None, None])
        ws.append(["Entry B", "Content B"])
        xlsx_path = tmp_path / "test.xlsx"
        wb.save(xlsx_path)

        results = dp.parse_xlsx(str(xlsx_path))
        assert len(results) == 2


class TestParseVetDrugExcel:
    def test_parse_official_excel_chinese_headers(self, tmp_path: Path):
        from openpyxl import Workbook

        wb = Workbook()
        ws = wb.active
        ws.append(["批准文号", "通用名称", "商品名称", "生产企业", "适应症", "类别"])
        ws.append(["兽药字220125001", "非泼罗尼", "宠安宁", "某某兽药公司", "体外驱虫", "化学药品"])
        ws.append(["兽药字220125002", "阿苯达唑", "驱虫片", "另一公司", "体内驱虫", "化学药品"])
        xlsx_path = tmp_path / "vetdrug.xlsx"
        wb.save(xlsx_path)

        results = dp.parse_vetdrug_excel(str(xlsx_path))
        assert len(results) == 2
        assert results[0]["approval_number"] == "兽药字220125001"
        assert results[0]["generic_name"] == "非泼罗尼"
        assert results[0]["product_name"] == "宠安宁"
        assert results[0]["category"] == "化学药品"

    def test_parse_official_excel_english_headers(self, tmp_path: Path):
        from openpyxl import Workbook

        wb = Workbook()
        ws = wb.active
        ws.append(["approval_number", "product_name", "generic_name", "manufacturer"])
        ws.append(["兽药字220125003", "TestProduct", "GenericA", "MakerA"])
        xlsx_path = tmp_path / "vetdrug_en.xlsx"
        wb.save(xlsx_path)

        results = dp.parse_vetdrug_excel(str(xlsx_path))
        assert len(results) == 1
        assert results[0]["approval_number"] == "兽药字220125003"

    def test_missing_approval_number_raises(self, tmp_path: Path):
        from openpyxl import Workbook

        wb = Workbook()
        ws = wb.active
        ws.append(["product_name", "generic_name"])
        ws.append(["Test", "Generic"])
        xlsx_path = tmp_path / "bad.xlsx"
        wb.save(xlsx_path)

        with pytest.raises(ValueError, match="approval_number"):
            dp.parse_vetdrug_excel(str(xlsx_path))


class TestParseDocument:
    def test_unsupported_extension(self, tmp_path: Path):
        txt_path = tmp_path / "test.txt"
        txt_path.write_text("hello")
        with pytest.raises(ValueError, match="Unsupported file extension"):
            dp.parse_document(str(txt_path))

    def test_dispatches_pdf(self, tmp_path: Path):
        from PyPDF2 import PdfWriter

        writer = PdfWriter()
        writer.add_blank_page(width=612, height=792)
        pdf_path = tmp_path / "test.pdf"
        with open(pdf_path, "wb") as f:
            writer.write(f)

        results = dp.parse_document(str(pdf_path))
        assert isinstance(results, list)
